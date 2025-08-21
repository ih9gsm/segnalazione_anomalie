from io import BytesIO
from typing import List, Optional, Tuple

from reportlab.lib.pagesizes import A4
from reportlab.pdfgen import canvas
from docx import Document
from docx.shared import Inches
from email.message import EmailMessage
import smtplib

from .models import SMTPSettings


def generate_pdf(description: str, logo_path: Optional[str] = None) -> bytes:
    buffer = BytesIO()
    c = canvas.Canvas(buffer, pagesize=A4)
    width, height = A4
    y = height - 50
    if logo_path:
        try:
            c.drawImage(logo_path, 40, y - 40, width=100, preserveAspectRatio=True, mask='auto')
            y -= 80
        except Exception:
            pass
    c.drawString(40, y, f"Report:")
    y -= 20
    text = c.beginText(40, y)
    for line in description.splitlines():
        text.textLine(line)
    c.drawText(text)
    c.showPage()
    c.save()
    buffer.seek(0)
    return buffer.read()


def generate_docx(description: str, logo_path: Optional[str] = None) -> bytes:
    doc = Document()
    if logo_path:
        try:
            doc.add_picture(logo_path, width=Inches(2))
        except Exception:
            pass
    doc.add_paragraph("Report:")
    for line in description.splitlines():
        doc.add_paragraph(line)
    buffer = BytesIO()
    doc.save(buffer)
    buffer.seek(0)
    return buffer.read()


def send_email(
    smtp: SMTPSettings,
    subject: str,
    body: str,
    recipients: List[str],
    attachments: Optional[List[Tuple[str, bytes]]] = None,
) -> None:
    msg = EmailMessage()
    msg["Subject"] = subject
    msg["From"] = smtp.username or "noreply@example.com"
    msg["To"] = ", ".join(recipients)
    msg.set_content(body)
    for filename, content in attachments or []:
        msg.add_attachment(content, maintype="application", subtype="octet-stream", filename=filename)

    with smtplib.SMTP(smtp.server, smtp.port) as server:
        if smtp.use_tls:
            server.starttls()
        if smtp.username and smtp.password:
            server.login(smtp.username, smtp.password)
        server.send_message(msg)
